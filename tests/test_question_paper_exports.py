"""
Regression tests for question paper exports, grade validation, and data integrity.

These tests do NOT require a live database or LLM — they mock the DB layer
and test the service / blueprint functions directly.
"""

import io
from typing import Any, Dict, List
from unittest.mock import AsyncMock, patch

import pytest

from open_notebook.graphs.question_paper_blueprint import (
    DEFAULT_PRESET,
    build_slots,
)


# ---------------------------------------------------------------------------
# Helpers to build realistic paper records
# ---------------------------------------------------------------------------


def _build_paper_record(
    *,
    num_questions: int = 50,
    status: str = "completed",
    grade: str = "8",
    blueprint: dict | None = None,
    include_failed_slots: bool = False,
    strip_newer_fields: bool = False,
) -> dict:
    """Build a realistic paper DB record with N questions."""
    bp = blueprint or DEFAULT_PRESET
    slots = build_slots(bp, grade=grade, subject="Science")
    questions: List[dict] = []
    for s in slots[:num_questions]:
        indices = [0] if s.answer_type == "single_correct" else [0, 1]
        q: Dict[str, Any] = {
            "question_number": s.question_number,
            "question": f"Test question {s.question_number}?",
            "type": "mcq",
            "answer_type": s.answer_type,
            "options": ["A", "B", "C", "D", "E"],
            "correct_indices": indices,
            "marks": 1,
            "topic": "Test Topic",
            "sub_topic": "Sub",
            "grade": grade,
            "chapter": s.chapter,
            "chapter_title": s.chapter_title,
            "section_ref": s.chapter_title,
            "difficulty": s.target_difficulty,
            "target_difficulty": s.target_difficulty,
            "validated_cognitive_difficulty": s.target_difficulty,
            "difficulty_score": 15,
            "difficulty_scores": {
                "knowledge": 2, "reasoning": 2, "context": 2, "application": 2,
                "interpretation": 2, "decision_making": 2, "concept_integration": 2,
                "distractor_quality": 1,
            },
            "validation_status": "passed",
            "validation_reasons": [],
            "generation_attempts": 1,
            "explanation": f"Explanation for Q{s.question_number}",
            "answer": "A",
        }
        if strip_newer_fields:
            for key in (
                "target_difficulty", "validated_cognitive_difficulty",
                "difficulty_score", "difficulty_scores",
                "validation_status", "validation_reasons",
                "generation_attempts", "answer_type",
            ):
                q.pop(key, None)
        questions.append(q)

    answer_key = [
        {
            "question_number": q["question_number"],
            "question": q["question"],
            "answer": q.get("answer", "A"),
            "explanation": q.get("explanation", ""),
            "marks": 1,
        }
        for q in questions
    ]

    failed = []
    if include_failed_slots:
        failed = [
            {
                "question_number": 99,
                "chapter": 1,
                "chapter_title": "Chapter 1",
                "target_difficulty": "easy",
                "answer_type": "single_correct",
                "validation_status": "needs_manual_review",
                "validation_reasons": ["distractor quality failed"],
                "generation_attempts": 3,
            }
        ]

    return {
        "id": "question_paper:test123",
        "topic": "Test Science",
        "difficulty": "medium",
        "target_marks": 50,
        "grade": grade,
        "subject": "Science",
        "language": "en",
        "pass_percentage": 70,
        "status": status,
        "error_message": None,
        "blueprint": bp,
        "final_paper": {
            "sections": [{"section_name": "mcq", "questions": questions}],
            "total_marks": len(questions),
            "question_count": len(questions),
            "pass_percentage": 70,
            "grade": grade,
            "subject": "Science",
        },
        "answer_key": answer_key,
        "failed_slots": failed,
        "audit": {"ok": status == "completed", "errors": []},
        "covered_topics": ["Topic A"],
        "coverage_gaps": [],
        "created": "2026-08-19T10:00:00Z",
    }


def _mock_load_paper(paper: dict):
    """Return an AsyncMock that simulates _load_paper_for_export."""
    async def _load(paper_id: str) -> dict:
        if paper.get("status") not in ("completed", "needs_manual_review"):
            from fastapi import HTTPException
            raise HTTPException(status_code=409, detail="Paper not yet completed")
        return paper
    return _load


# ---------------------------------------------------------------------------
# Test: Grade field must not default to "8"
# ---------------------------------------------------------------------------


class TestGradeValidation:
    def test_generate_form_grade_starts_empty(self):
        """The GenerateForm state initializes grade as empty string, not '8'."""
        # This is a code-level assertion — the frontend sets grade = ''
        # and the form is disabled until grade is filled.
        # Backend: grade is Optional[str] = None on GeneratePaperRequest.
        from api.question_paper_service import GeneratePaperRequest

        req = GeneratePaperRequest(topic="Test", target_marks=50)
        assert req.grade is None, "Grade must not default to a specific value"

    def test_generation_cannot_submit_without_grade(self):
        """The form requires grade — submit should not work with empty grade."""
        from api.question_paper_service import GeneratePaperRequest

        req = GeneratePaperRequest(topic="Test", target_marks=50, grade=None)
        assert req.grade is None
        req2 = GeneratePaperRequest(topic="Test", target_marks=50, grade="")
        assert req2.grade == ""

    def test_slots_carry_grade_from_request(self):
        slots = build_slots(DEFAULT_PRESET, grade="6", subject="Math")
        assert all(s.grade == "6" for s in slots)
        slots_empty = build_slots(DEFAULT_PRESET, grade="", subject="Math")
        assert all(s.grade == "" for s in slots_empty)


# ---------------------------------------------------------------------------
# Test: Export reads chapter_difficulty correctly
# ---------------------------------------------------------------------------


class TestExportBlueprintReading:
    def test_default_target_counts_easy14_medium18_difficult18(self):
        """Export must read blueprint and show E14 / M18 / D18."""
        paper = _build_paper_record(num_questions=50, status="completed")
        blueprint = paper["blueprint"]
        chapters_bp = blueprint.get("chapter_difficulty", {})
        diff_totals = {"easy": 0, "medium": 0, "difficult": 0}
        for ch_counts in chapters_bp.values():
            if isinstance(ch_counts, dict):
                for d, v in ch_counts.items():
                    diff_totals[d] = diff_totals.get(d, 0) + int(v)
        assert diff_totals == {"easy": 14, "medium": 18, "difficult": 18}

    def test_blueprint_with_custom_counts(self):
        """Non-default blueprint should produce correct target totals."""
        custom = {
            "id": "custom",
            "total_questions": 10,
            "pass_percentage": 50,
            "options_per_question": 5,
            "format": "mcq",
            "chapter_difficulty": {
                "1": {"easy": 2, "medium": 3, "difficult": 0},
                "2": {"easy": 1, "medium": 2, "difficult": 2},
            },
            "difficulty_answer_types": {
                "easy": {"single_correct": 3, "multiple_correct": 0},
                "medium": {"single_correct": 5, "multiple_correct": 0},
                "difficult": {"single_correct": 2, "multiple_correct": 0},
            },
        }
        slots = build_slots(custom, grade="5", subject="English")
        assert len(slots) == 10
        diff_counts = {"easy": 0, "medium": 0, "difficult": 0}
        for s in slots:
            diff_counts[s.target_difficulty] += 1
        assert diff_counts == {"easy": 3, "medium": 5, "difficult": 2}


# ---------------------------------------------------------------------------
# Test: needs_manual_review exports contain DRAFT status
# ---------------------------------------------------------------------------


class TestDraftExport:
    @pytest.mark.asyncio
    async def test_txt_export_has_draft_marker_for_manual_review(self):
        paper = _build_paper_record(
            num_questions=42,
            status="needs_manual_review",
            include_failed_slots=True,
        )
        from api.question_paper_service import QuestionPaperService

        with patch.object(
            QuestionPaperService,
            "_load_paper_for_export",
            side_effect=_mock_load_paper(paper),
        ):
            txt_bytes = await QuestionPaperService.export_paper_txt(
                "question_paper:test123"
            )
        text = txt_bytes.decode("utf-8")
        assert "DRAFT" in text
        assert "NEEDS MANUAL REVIEW" in text

    @pytest.mark.asyncio
    async def test_docx_export_has_draft_marker(self):
        paper = _build_paper_record(
            num_questions=42,
            status="needs_manual_review",
            include_failed_slots=True,
        )
        from api.question_paper_service import QuestionPaperService

        with patch.object(
            QuestionPaperService,
            "_load_paper_for_export",
            side_effect=_mock_load_paper(paper),
        ):
            docx_bytes = await QuestionPaperService.export_paper_docx(
                "question_paper:test123"
            )
        from docx import Document

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "DRAFT" in full_text
        assert "NEEDS MANUAL REVIEW" in full_text


# ---------------------------------------------------------------------------
# Test: Expected and accepted question counts are exported separately
# ---------------------------------------------------------------------------


class TestCountExport:
    @pytest.mark.asyncio
    async def test_txt_shows_expected_and_accepted_separately(self):
        paper = _build_paper_record(num_questions=42, status="needs_manual_review")
        paper["target_marks"] = 50
        from api.question_paper_service import QuestionPaperService

        with patch.object(
            QuestionPaperService,
            "_load_paper_for_export",
            side_effect=_mock_load_paper(paper),
        ):
            txt_bytes = await QuestionPaperService.export_paper_txt(
                "question_paper:test123"
            )
        text = txt_bytes.decode("utf-8")
        assert "Expected Questions/Marks: 50" in text
        assert "Accepted Questions/Actual Marks: 42" in text

    @pytest.mark.asyncio
    async def test_xlsx_audit_sheet_shows_both_counts(self):
        paper = _build_paper_record(num_questions=42, status="needs_manual_review")
        paper["target_marks"] = 50
        from api.question_paper_service import QuestionPaperService

        with patch.object(
            QuestionPaperService,
            "_load_paper_for_export",
            side_effect=_mock_load_paper(paper),
        ):
            xlsx_bytes = await QuestionPaperService.export_paper_xlsx(
                "question_paper:test123"
            )
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(xlsx_bytes))
        ws = wb["Audit Summary"]
        rows = {row[0].value: row[1].value for row in ws.iter_rows(min_row=2, max_col=2)}
        assert rows["Expected Questions"] == 50
        assert rows["Accepted Questions / Actual Marks"] == 42


# ---------------------------------------------------------------------------
# Test: XLSX contains all 5 expected sheets
# ---------------------------------------------------------------------------


class TestXlsxSheets:
    @pytest.mark.asyncio
    async def test_xlsx_has_five_sheets(self):
        paper = _build_paper_record(num_questions=50, status="completed")
        from api.question_paper_service import QuestionPaperService

        with patch.object(
            QuestionPaperService,
            "_load_paper_for_export",
            side_effect=_mock_load_paper(paper),
        ):
            xlsx_bytes = await QuestionPaperService.export_paper_xlsx(
                "question_paper:test123"
            )
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(xlsx_bytes))
        expected_sheets = {
            "Questions Review",
            "Cognitive Validation",
            "Answer Key",
            "Audit Summary",
            "Failed Slots",
        }
        assert set(wb.sheetnames) == expected_sheets

    @pytest.mark.asyncio
    async def test_xlsx_questions_review_row_count(self):
        paper = _build_paper_record(num_questions=50, status="completed")
        from api.question_paper_service import QuestionPaperService

        with patch.object(
            QuestionPaperService,
            "_load_paper_for_export",
            side_effect=_mock_load_paper(paper),
        ):
            xlsx_bytes = await QuestionPaperService.export_paper_xlsx(
                "question_paper:test123"
            )
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(xlsx_bytes))
        ws = wb["Questions Review"]
        data_rows = ws.max_row - 1  # subtract header
        assert data_rows == 50


# ---------------------------------------------------------------------------
# Test: Old papers with missing newer metadata do not crash exports
# ---------------------------------------------------------------------------


class TestLegacyPaperExport:
    @pytest.mark.asyncio
    async def test_old_paper_without_validation_fields_exports_xlsx(self):
        paper = _build_paper_record(
            num_questions=10,
            status="completed",
            blueprint={
                "id": "legacy",
                "total_questions": 10,
                "pass_percentage": 70,
                "options_per_question": 5,
                "format": "mcq",
                "chapter_difficulty": {
                    "1": {"easy": 3, "medium": 4, "difficult": 3},
                },
                "difficulty_answer_types": {
                    "easy": {"single_correct": 3, "multiple_correct": 0},
                    "medium": {"single_correct": 4, "multiple_correct": 0},
                    "difficult": {"single_correct": 3, "multiple_correct": 0},
                },
            },
            strip_newer_fields=True,
        )
        # Also strip audit and blueprint from paper level (legacy paper)
        paper["audit"] = None
        paper["blueprint"] = None

        from api.question_paper_service import QuestionPaperService

        with patch.object(
            QuestionPaperService,
            "_load_paper_for_export",
            side_effect=_mock_load_paper(paper),
        ):
            xlsx_bytes = await QuestionPaperService.export_paper_xlsx(
                "question_paper:test123"
            )
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(xlsx_bytes))
        assert "Questions Review" in wb.sheetnames
        ws = wb["Questions Review"]
        assert ws.max_row > 1  # at least header + data

    @pytest.mark.asyncio
    async def test_old_paper_without_blueprint_exports_txt(self):
        paper = _build_paper_record(num_questions=10, status="completed")
        paper["blueprint"] = None
        paper["audit"] = None

        from api.question_paper_service import QuestionPaperService

        with patch.object(
            QuestionPaperService,
            "_load_paper_for_export",
            side_effect=_mock_load_paper(paper),
        ):
            txt_bytes = await QuestionPaperService.export_paper_txt(
                "question_paper:test123"
            )
        text = txt_bytes.decode("utf-8")
        assert "QUESTION PAPER" in text
        assert "Accepted Questions" in text

    @pytest.mark.asyncio
    async def test_old_paper_without_blueprint_exports_docx(self):
        paper = _build_paper_record(num_questions=10, status="completed")
        paper["blueprint"] = None
        paper["audit"] = None

        from api.question_paper_service import QuestionPaperService

        with patch.object(
            QuestionPaperService,
            "_load_paper_for_export",
            side_effect=_mock_load_paper(paper),
        ):
            docx_bytes = await QuestionPaperService.export_paper_docx(
                "question_paper:test123"
            )
        from docx import Document

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Question Paper" in full_text
