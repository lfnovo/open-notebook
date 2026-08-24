"""
Service layer for question paper generation.
Handles paper creation, job submission, status tracking, and book upload.
"""

import csv
import json
import io
import os
import re
from typing import Any, Dict, List, Optional

from fastapi import HTTPException, UploadFile
from loguru import logger
from pydantic import BaseModel
from surreal_commands import get_command_status, submit_command

from open_notebook.database.repository import repo_create, repo_delete, repo_query, repo_update
from open_notebook.domain.question_bank import QuestionPaper, QuestionRecord

# Register command module at import time so it's always available
try:
    import commands.question_paper_commands  # noqa: F401
except ImportError as e:
    logger.warning(f"Could not import question paper commands module: {e}")


# ---------------------------------------------------------------------------
# Request / Response models
# ---------------------------------------------------------------------------

class GeneratePaperRequest(BaseModel):
    topic: str
    difficulty: str = "medium"  # legacy overall hint; generation uses the blueprint
    target_marks: int = 50
    section_config: dict = {}
    curriculum_objectives: List[str] = []
    generator_model: Optional[str] = None
    reviewer_model: Optional[str] = None
    book_id: Optional[str] = None
    selected_chapters: Optional[List[int]] = None
    skip_front_matter: bool = True
    grade: Optional[str] = None
    subject: Optional[str] = None
    language: str = "en"
    pass_percentage: int = 70
    options_per_question: int = 5
    question_format: str = "mcq"
    blueprint: Optional[Dict[str, Any]] = None
    max_slot_attempts: int = 3
    slot_concurrency: int = 3


class GeneratePaperResponse(BaseModel):
    job_id: str
    paper_id: str
    status: str
    message: str
    topic: str


class GenerateBankBatchRequest(BaseModel):
    book_id: str
    grade: str
    subject: str
    chapter: int
    difficulty: str
    total_questions: int
    single_correct: int
    multiple_correct: int
    language: str = "en"
    curriculum_objectives: List[str] = []
    generator_model: Optional[str] = None
    reviewer_model: Optional[str] = None
    max_slot_attempts: int = 3
    max_refill_attempts: int = 5
    slot_concurrency: int = 3


class GenerateBankBatchResponse(BaseModel):
    job_id: str
    batch_id: str
    status: str
    message: str
    requested: int


def _as_int(value: Any, default: int = 0) -> int:
    try:
        if value is None:
            return default
        return int(value)
    except (TypeError, ValueError):
        return default


def _section_config_total(section_config: Any) -> int:
    if not isinstance(section_config, dict):
        return 0
    total = 0
    for value in section_config.values():
        try:
            total += int(value or 0)
        except (TypeError, ValueError):
            continue
    return total


def _requested_question_count(record: Dict[str, Any]) -> tuple[Optional[int], str]:
    """Requested question total from persisted question-count fields only.

    Never uses target_marks (marks, not questions). Returns (None, "unknown")
    when no requested-question count is stored.
    """
    blueprint = record.get("blueprint")
    if isinstance(blueprint, dict) and blueprint.get("total_questions") is not None:
        total = _as_int(blueprint.get("total_questions"))
        if total > 0:
            return total, "blueprint.total_questions"
    nested_total = record.get("blueprint.total_questions")
    if nested_total is not None:
        total = _as_int(nested_total)
        if total > 0:
            return total, "blueprint.total_questions"
    section_total = _section_config_total(record.get("section_config"))
    if section_total > 0:
        return section_total, "section_config"
    for key in ("requested_questions", "requested"):
        if record.get(key) is not None:
            total = _as_int(record.get(key))
            if total > 0:
                return total, key
    return None, "unknown"


def _generated_question_count(record: Dict[str, Any], *, section_count: Optional[int] = None) -> int:
    """Accepted/generated count from persisted final_paper. Does not change stored records."""
    if section_count is not None:
        return _as_int(section_count)
    for key in ("generated_questions", "final_paper.question_count"):
        if record.get(key) is not None:
            return _as_int(record.get(key))
    final_paper = record.get("final_paper")
    if isinstance(final_paper, dict) and final_paper.get("question_count") is not None:
        return _as_int(final_paper.get("question_count"))
    return 0


def _history_display_status(
    stored_status: Any,
    generated: int,
    requested: Optional[int],
) -> str:
    """Display-only status. Stored DB status is not rewritten."""
    status = str(stored_status or "").lower()
    if status in ("running", "pending"):
        return "running"
    if requested is not None and requested > 0 and generated == requested:
        return "completed"
    if requested is not None and generated > 0 and generated < requested:
        return "partial"
    if generated == 0 and status in ("failed", "error"):
        return "failed"
    if status in ("needs_manual_review", "needs_review", "needs_manual_review"):
        return "partial"
    return status


def _question_progress_fields(
    record: Dict[str, Any],
    *,
    section_count: Optional[int] = None,
) -> Dict[str, Any]:
    requested, requested_source = _requested_question_count(record)
    generated = _generated_question_count(record, section_count=section_count)
    remaining = max(0, requested - generated) if requested is not None else None
    stored_status = record.get("status") or ""
    mix_fields = _difficulty_mix_fields(record)
    return {
        "requested_questions": requested,
        "generated_questions": generated,
        "remaining_questions": remaining,
        "requested_source": requested_source,
        "display_status": _history_display_status(stored_status, generated, requested),
        **mix_fields,
    }


_DIFFICULTY_KEYS = ("easy", "medium", "difficult")
_DIFFICULTY_LABELS = {"easy": "Easy", "medium": "Medium", "difficult": "Difficult"}


def _empty_difficulty_mix() -> Dict[str, int]:
    return {"easy": 0, "medium": 0, "difficult": 0}


def _normalize_difficulty(value: Any) -> Optional[str]:
    raw = str(value or "").strip().lower()
    if raw in ("hard", "difficult"):
        return "difficult"
    if raw in ("easy", "medium"):
        return raw
    return None


def _add_count(mix: Dict[str, int], key: Any, amount: int) -> None:
    normalized = _normalize_difficulty(key)
    if normalized and amount:
        mix[normalized] += amount


def _requested_difficulty_mix(record: Dict[str, Any]) -> Optional[Dict[str, int]]:
    """Requested Easy/Medium/Difficult totals from persisted blueprint only."""
    blueprint = record.get("blueprint")
    if not isinstance(blueprint, dict):
        return None
    mix = _empty_difficulty_mix()
    found = False
    matrix = blueprint.get("chapter_difficulty")
    if isinstance(matrix, dict):
        for counts in matrix.values():
            if not isinstance(counts, dict):
                continue
            for key, value in counts.items():
                n = _as_int(value)
                if n:
                    _add_count(mix, key, n)
                    found = True
    if found and sum(mix.values()) > 0:
        return mix
    answer_types = blueprint.get("difficulty_answer_types")
    if isinstance(answer_types, dict):
        for diff, counts in answer_types.items():
            if not isinstance(counts, dict):
                continue
            n = _as_int(counts.get("single_correct")) + _as_int(counts.get("multiple_correct"))
            if n:
                _add_count(mix, diff, n)
                found = True
    if found and sum(mix.values()) > 0:
        return mix
    return None


def _generated_difficulty_mix(final_paper: Any) -> Dict[str, int]:
    """Actual accepted difficulty counts from persisted final_paper questions."""
    mix = _empty_difficulty_mix()
    if not isinstance(final_paper, dict):
        return mix
    for section in final_paper.get("sections") or []:
        if not isinstance(section, dict):
            continue
        for question in section.get("questions") or []:
            if not isinstance(question, dict):
                continue
            difficulty = _normalize_difficulty(
                question.get("validated_cognitive_difficulty")
                or question.get("target_difficulty")
                or question.get("difficulty")
            )
            if difficulty:
                mix[difficulty] += 1
    return mix


def _difficulty_mix_kind(mix: Optional[Dict[str, int]]) -> str:
    if not mix:
        return "unknown"
    present = [key for key in _DIFFICULTY_KEYS if mix.get(key, 0) > 0]
    if len(present) == 1:
        return f"{present[0]}_only"
    if len(present) > 1:
        return "mixed"
    return "unknown"


def _difficulty_mix_label(mix: Optional[Dict[str, int]]) -> str:
    if not mix:
        return "—"
    present = [(key, mix[key]) for key in _DIFFICULTY_KEYS if mix.get(key, 0) > 0]
    if not present:
        return "—"
    if len(present) == 1:
        key, count = present[0]
        return f"{_DIFFICULTY_LABELS[key]} Only · {count}"
    return " · ".join(f"{_DIFFICULTY_LABELS[key]} {count}" for key, count in present)


def _difficulty_mix_fields(
    record: Dict[str, Any],
    *,
    final_paper: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    requested_mix = _requested_difficulty_mix(record)
    paper = final_paper if final_paper is not None else record.get("final_paper")
    generated_mix = _generated_difficulty_mix(paper) if paper else None
    remaining_mix = None
    if requested_mix is not None and generated_mix is not None:
        remaining_mix = {
            key: max(0, requested_mix.get(key, 0) - generated_mix.get(key, 0))
            for key in _DIFFICULTY_KEYS
        }
    return {
        "requested_difficulty": requested_mix,
        "generated_difficulty": generated_mix,
        "remaining_difficulty": remaining_mix,
        "difficulty_mix": _difficulty_mix_kind(requested_mix),
        "difficulty_mix_label": _difficulty_mix_label(requested_mix),
    }


# ---------------------------------------------------------------------------
# Service
# ---------------------------------------------------------------------------

class QuestionPaperService:

    @staticmethod
    async def create_and_submit(request: GeneratePaperRequest) -> GeneratePaperResponse:
        """Create a question_paper record and submit the generation job."""

        from open_notebook.graphs.question_paper_blueprint import (
            build_slots,
            preset_from_dict,
            validate_chapter_selection,
        )

        valid_difficulties = {"easy", "medium", "hard", "difficult", "mixed"}
        if request.difficulty not in valid_difficulties:
            raise HTTPException(
                status_code=400,
                detail=f"difficulty must be one of {sorted(valid_difficulties)}",
            )
        if request.target_marks < 1 or request.target_marks > 500:
            raise HTTPException(
                status_code=400,
                detail="target_marks must be between 1 and 500",
            )
        if request.section_config and any(
            not isinstance(v, int) or v < 1 or v > 100 for v in request.section_config.values()
        ):
            raise HTTPException(status_code=400, detail="section_config values must be positive integers")
        if not request.topic.strip():
            raise HTTPException(status_code=400, detail="topic cannot be empty")
        if request.book_id and request.selected_chapters is not None and len(request.selected_chapters) == 0:
            raise HTTPException(status_code=400, detail="selected_chapters cannot be empty; omit it to use all chapters")
        if request.options_per_question != 5:
            raise HTTPException(status_code=400, detail="This workflow requires exactly 5 options per question")
        if request.max_slot_attempts < 1 or request.max_slot_attempts > 10:
            raise HTTPException(status_code=400, detail="max_slot_attempts must be between 1 and 10")
        if request.slot_concurrency < 1 or request.slot_concurrency > 8:
            raise HTTPException(status_code=400, detail="slot_concurrency must be between 1 and 8")

        if not request.blueprint or not request.blueprint.get("difficulty_answer_types"):
            raise HTTPException(
                status_code=400,
                detail="difficulty_answer_types is required — manually configure Single Correct and Multiple Correct counts per difficulty",
            )
        bp = request.blueprint
        if bp.get("chapter_difficulty") and bp.get("total_questions") is not None:
            matrix_sum = sum(
                sum(int(v) for v in counts.values())
                for counts in bp["chapter_difficulty"].values()
            )
            if matrix_sum != int(bp["total_questions"]):
                raise HTTPException(
                    status_code=400,
                    detail=f"Chapter matrix totals {matrix_sum} questions but total_questions is {bp['total_questions']}. They must match.",
                )
        preset = preset_from_dict(request.blueprint)
        try:
            slots = build_slots(
                preset,
                grade=request.grade or "",
                subject=request.subject or request.topic.strip(),
            )
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))
        if len(slots) != preset["total_questions"]:
            raise HTTPException(
                status_code=400,
                detail="Blueprint chapter and answer-type matrices do not produce a consistent slot count",
            )

        book_content: Optional[str] = None
        book_chapters: List[Dict[str, Any]] = []
        if request.book_id:
            book_chapters = await QuestionPaperService.get_book_chapters(
                request.book_id, request.selected_chapters,
                skip_front_matter=request.skip_front_matter,
            )
            book_content = await QuestionPaperService.get_book_content(
                request.book_id, request.selected_chapters,
                skip_front_matter=request.skip_front_matter,
            )
            if not book_chapters:
                raise HTTPException(
                    status_code=400,
                    detail="No chapter content could be loaded from the selected learning material.",
                )
            try:
                validate_chapter_selection(preset, len(book_chapters))
            except ValueError as e:
                raise HTTPException(status_code=400, detail=str(e))

        derived_section_config = request.section_config or {
            "mcq": sum(1 for s in slots if s.answer_type == "single_correct"),
            "multi_correct": sum(1 for s in slots if s.answer_type == "multiple_correct"),
        }

        paper_data = {
            "topic": request.topic.strip(),
            "difficulty": request.difficulty,
            "target_marks": request.target_marks,
            "section_config": derived_section_config,
            "status": "pending",
            "grade": request.grade,
            "subject": (request.subject or request.topic).strip(),
            "language": request.language,
            "pass_percentage": request.pass_percentage,
            "blueprint": preset,
        }
        raw_paper = await repo_create("question_paper", paper_data)
        if not raw_paper:
            raise HTTPException(status_code=500, detail="Failed to create paper record")
        paper_record = raw_paper[0] if isinstance(raw_paper, list) else raw_paper

        paper_id = str(paper_record["id"])

        try:
            cmd_id = submit_command(
                "open_notebook",
                "generate_question_paper",
                {
                    "paper_id": paper_id,
                    "topic": request.topic.strip(),
                    "difficulty": request.difficulty,
                    "target_marks": request.target_marks,
                    "section_config": derived_section_config,
                    "curriculum_objectives": request.curriculum_objectives,
                    "generator_model": request.generator_model,
                    "reviewer_model": request.reviewer_model,
                    "book_content": book_content,
                    "book_chapters": book_chapters,
                    "grade": request.grade,
                    "subject": (request.subject or request.topic).strip(),
                    "language": request.language,
                    "pass_percentage": request.pass_percentage,
                    "blueprint_preset": preset,
                    "max_slot_attempts": request.max_slot_attempts,
                    "slot_concurrency": request.slot_concurrency,
                },
            )
        except Exception as e:
            logger.error(f"Failed to submit question paper job: {e}")
            raise HTTPException(status_code=500, detail="Failed to submit generation job")

        cmd_id_str = str(cmd_id)

        # 6. Store command reference on paper record using raw SurrealQL
        # The command field is option<record> — must use a record link literal in SurrealQL,
        # not a plain string or RecordID object passed via MERGE.
        from open_notebook.database.repository import ensure_record_id, repo_query
        try:
            await repo_query(
                f"UPDATE {paper_id} SET command = {cmd_id_str}",
                {},
            )
        except Exception as e:
            # Non-fatal: job was already submitted; log and continue
            logger.warning(f"Could not store command reference on paper {paper_id}: {e}")

        logger.info(f"Submitted question paper job {cmd_id_str} for paper {paper_id}")

        return GeneratePaperResponse(
            job_id=cmd_id_str,
            paper_id=paper_id,
            status="submitted",
            message=f"Question paper generation started for topic '{request.topic}'",
            topic=request.topic,
        )

    @staticmethod
    async def get_paper_status(paper_id: str) -> Dict[str, Any]:
        """Get status of a question paper by its record ID."""
        from open_notebook.database.repository import ensure_record_id
        results = await repo_query(
            "SELECT * FROM question_paper WHERE id = $id",
            {"id": ensure_record_id(paper_id)},
        )
        if not results:
            raise HTTPException(status_code=404, detail="Paper not found")

        paper = results[0]
        paper_status = paper.get("status", "unknown")
        job_status = None

        command_id = paper.get("command")
        if command_id:
            try:
                cmd = await get_command_status(str(command_id))
                if cmd:
                    job_status = cmd.status
                    # Sync failed status from command back to DB if not already recorded
                    if cmd.status in ("failed", "error") and paper_status not in (
                        "failed", "error", "needs_manual_review", "completed",
                    ):
                        paper_status = "failed"
                        from open_notebook.database.repository import repo_update
                        try:
                            await repo_update(
                                "question_paper",
                                ensure_record_id(paper_id),
                                {"status": "failed", "error_message": "Job execution failed"},
                            )
                        except Exception:
                            pass
            except Exception as e:
                logger.warning(f"Failed to get command status for {command_id}: {e}")

        return {
            "paper_id": str(paper.get("id", paper_id)),
            "status": paper_status,
            "job_status": job_status,
            "topic": paper.get("topic", ""),
            "difficulty": paper.get("difficulty", ""),
            "target_marks": paper.get("target_marks", 0),
            "error_message": paper.get("error_message"),
            "grade": paper.get("grade"),
            "created": str(paper.get("created", "")),
        }

    @staticmethod
    async def get_paper_result(paper_id: str) -> Dict[str, Any]:
        """Fetch a completed paper result."""
        from open_notebook.database.repository import ensure_record_id
        results = await repo_query(
            "SELECT * FROM question_paper WHERE id = $id",
            {"id": ensure_record_id(paper_id)},
        )
        if not results:
            raise HTTPException(status_code=404, detail="Paper not found")

        paper = results[0]
        if paper.get("status") not in ("completed", "needs_manual_review", "needs_manual_review"):
            raise HTTPException(
                status_code=409,
                detail=f"Paper is not yet completed (status: {paper.get('status')})",
            )

        actual_id = str(paper.get("id", paper_id))
        final_paper = paper.get("final_paper") or {}
        section_count = sum(
            len(s.get("questions", [])) for s in final_paper.get("sections", [])
        )
        stored_count = final_paper.get("question_count", 0)
        if stored_count != section_count:
            logger.warning(
                f"Paper {actual_id}: question_count field ({stored_count}) differs "
                f"from section question count ({section_count}). Using section count."
            )
            final_paper["question_count"] = section_count
        logger.info(
            f"Result API paper {actual_id}: "
            f"status={paper.get('status')}, accepted_questions={section_count}"
        )

        progress = _question_progress_fields(paper, section_count=section_count)

        return {
            "paper_id": actual_id,
            "topic": paper.get("topic", ""),
            "difficulty": paper.get("difficulty", ""),
            "target_marks": paper.get("target_marks", 0),
            "section_config": paper.get("section_config", {}),
            "final_paper": final_paper,
            "answer_key": paper.get("answer_key", []),
            "coverage_gaps": paper.get("coverage_gaps", []),
            "covered_topics": paper.get("covered_topics", []),
            "status": paper.get("status", ""),
            "grade": paper.get("grade"),
            "subject": paper.get("subject"),
            "language": paper.get("language"),
            "pass_percentage": paper.get("pass_percentage"),
            "blueprint": paper.get("blueprint"),
            "audit": paper.get("audit"),
            "failed_slots": paper.get("failed_slots") or final_paper.get("failed_slots") or [],
            "error_message": paper.get("error_message"),
            "created": str(paper.get("created", "")),
            **progress,
        }

    @staticmethod
    async def list_papers() -> List[Dict[str, Any]]:
        """List all question papers (metadata only, no content)."""
        query = (
            "SELECT id, topic, difficulty, target_marks, section_config, status, "
            "error_message, grade, book_id, created, blueprint, "
            "final_paper.question_count AS generated_questions "
            "FROM question_paper ORDER BY created DESC LIMIT 50"
        )
        try:
            results = await repo_query(query)
        except Exception as e:
            logger.warning(f"Paper list nested count query failed, retrying without it: {e}")
            results = await repo_query(
                "SELECT id, topic, difficulty, target_marks, section_config, status, "
                "error_message, grade, book_id, created, blueprint "
                "FROM question_paper ORDER BY created DESC LIMIT 50"
            )

        papers = []
        for r in results or []:
            progress = _question_progress_fields(r)
            papers.append(
                {
                    "paper_id": str(r.get("id", "")),
                    "topic": r.get("topic", ""),
                    "difficulty": r.get("difficulty", ""),
                    "target_marks": r.get("target_marks", 0),
                    "section_config": r.get("section_config", {}),
                    "status": r.get("status", ""),
                    "error_message": r.get("error_message"),
                    "grade": r.get("grade"),
                    "book_id": str(r["book_id"]) if r.get("book_id") else None,
                    "question_count": progress["generated_questions"],
                    "created": str(r.get("created", "")),
                    **progress,
                }
            )
        return papers

    @staticmethod
    async def delete_paper(paper_id: str) -> None:
        """Delete a question paper record."""
        from open_notebook.database.repository import ensure_record_id
        try:
            await repo_delete(ensure_record_id(paper_id))
        except Exception as e:
            logger.error(f"Failed to delete paper {paper_id}: {e}")
            raise HTTPException(status_code=500, detail="Failed to delete paper")

    @staticmethod
    async def regenerate_missing(paper_id: str) -> Dict[str, Any]:
        """Re-run refill_slots on only the failed slots of an existing paper."""
        from open_notebook.database.repository import ensure_record_id
        from open_notebook.graphs.question_paper import (
            _slot_from_dict,
            audit_assembled_paper,
            assemble_paper,
            refill_slots,
            report_coverage,
        )
        from open_notebook.graphs.question_paper_blueprint import (
            MAX_REFILL_ATTEMPTS,
            MAX_SLOT_ATTEMPTS,
            audit_paper,
            build_effective_preset,
        )

        results = await repo_query(
            "SELECT * FROM question_paper WHERE id = $id",
            {"id": ensure_record_id(paper_id)},
        )
        if not results:
            raise HTTPException(status_code=404, detail="Paper not found")

        paper = results[0]
        if paper.get("status") != "needs_manual_review":
            raise HTTPException(
                status_code=409,
                detail=f"Paper status is '{paper.get('status')}'; only 'needs_manual_review' papers can be regenerated",
            )

        failed_slots = paper.get("failed_slots") or []
        if not failed_slots:
            raise HTTPException(status_code=409, detail="No failed slots to regenerate")

        final_paper = paper.get("final_paper") or {}
        approved = []
        for section in final_paper.get("sections", []):
            approved.extend(section.get("questions", []))

        # Build chapter_chunks from the stored book content if available
        chapter_chunks: Dict[str, List[str]] = {}
        book_chapters = paper.get("book_chapters") or []
        if not book_chapters:
            # Try to reconstruct from book_id
            book_id = paper.get("book_id")
            if book_id:
                try:
                    book_chapters = await QuestionPaperService.get_book_chapters(
                        book_id, paper.get("selected_chapters")
                    )
                except Exception:
                    pass

        if book_chapters:
            from open_notebook.graphs.question_paper_blueprint import chunk_chapter_text
            for i, ch in enumerate(book_chapters, start=1):
                text = ch.get("text", "")
                if text:
                    chapter_chunks[str(i)] = chunk_chapter_text(text)

        preset = paper.get("blueprint") or build_effective_preset(None)

        # Build state for refill_slots
        state = {
            "approved": approved,
            "failed_slots": failed_slots,
            "used_stems": [],
            "rejected_with_feedback": [],
            "chapter_chunks": chapter_chunks,
            "book_grounded": bool(chapter_chunks),
            "max_refill_attempts": MAX_REFILL_ATTEMPTS,
            "max_slot_attempts": MAX_SLOT_ATTEMPTS,
            "slot_concurrency": 3,
            "grade": paper.get("grade", ""),
            "subject": paper.get("subject", ""),
            "topic": paper.get("topic", ""),
            "language": paper.get("language", "en"),
            "generator_model": None,
            "reviewer_model": None,
            "effective_preset": preset,
            "pass_percentage": paper.get("pass_percentage", 70),
            "target_marks": paper.get("target_marks", len(approved) + len(failed_slots)),
        }

        logger.info(
            f"Regenerating {len(failed_slots)} missing slot(s) for paper {paper_id}"
        )

        refill_result = await refill_slots(state)

        # Merge refill results back
        merged_approved = refill_result.get("approved", approved)
        merged_failed = refill_result.get("failed_slots", [])

        # Rebuild state for assemble + audit
        assemble_state = {
            **state,
            "approved": merged_approved,
            "failed_slots": merged_failed,
        }
        assembled = await assemble_paper(assemble_state)
        assemble_state.update(assembled)

        coverage_result = await report_coverage(assemble_state)
        assemble_state.update(coverage_result)

        audit_result = await audit_assembled_paper(assemble_state)
        assemble_state.update(audit_result)

        # Persist
        new_final_paper = assemble_state.get("final_paper", {})
        new_answer_key = assemble_state.get("answer_key", [])
        audit = assemble_state.get("audit", {})
        audit_ok = bool(audit.get("ok"))
        new_status = "completed" if audit_ok else "needs_manual_review"
        error_message = None if audit_ok else "; ".join(audit.get("errors") or ["Audit failed"])

        import json as _json
        _fp_json = _json.dumps(new_final_paper)
        _ak_json = _json.dumps(new_answer_key)
        _cg_json = _json.dumps(assemble_state.get("coverage_gaps", []))
        _ct_json = _json.dumps(assemble_state.get("covered_topics", []))
        _fs_json = _json.dumps(merged_failed)
        _audit_json = _json.dumps(audit)
        _err_sql = "NONE" if error_message is None else _json.dumps(error_message)

        await repo_query(
            f"UPDATE {paper_id} MERGE "
            f"{{final_paper: {_fp_json}, answer_key: {_ak_json}, "
            f"coverage_gaps: {_cg_json}, covered_topics: {_ct_json}, "
            f"failed_slots: {_fs_json}, audit: {_audit_json}, "
            f"error_message: {_err_sql}, status: '{new_status}'}}",
            {},
        )

        accepted_count = new_final_paper.get("question_count", 0)
        logger.info(
            f"Regeneration complete for {paper_id}: "
            f"{accepted_count} accepted, {len(merged_failed)} still failed, status={new_status}"
        )

        return {
            "paper_id": str(paper.get("id", paper_id)),
            "status": new_status,
            "accepted_count": accepted_count,
            "still_failed": len(merged_failed),
            "audit": audit,
        }

    @staticmethod
    def _bank_search_item(r: Dict[str, Any]) -> Dict[str, Any]:
        """Map a stored question_bank row for the review UI (read-only)."""
        def _id(value: Any) -> Optional[str]:
            if value is None or value == "":
                return None
            return str(value)

        return {
            "id": str(r.get("id", "")),
            "question": r.get("question", ""),
            "topic": r.get("topic", ""),
            "sub_topic": r.get("sub_topic"),
            "type": r.get("type", ""),
            "difficulty": r.get("difficulty", ""),
            "answer": r.get("answer", ""),
            "explanation": r.get("explanation"),
            "options": r.get("options"),
            "correct_indices": r.get("correct_indices"),
            "grade": r.get("grade"),
            "subject": r.get("subject"),
            "chapter": r.get("chapter"),
            "chapter_title": r.get("chapter_title"),
            "answer_type": r.get("answer_type"),
            "target_difficulty": r.get("target_difficulty"),
            "validated_cognitive_difficulty": r.get("validated_cognitive_difficulty"),
            "difficulty_score": r.get("difficulty_score"),
            "validation_status": r.get("validation_status"),
            "batch_id": _id(r.get("batch_id")),
            "book_id": _id(r.get("book_id")),
        }

    @staticmethod
    async def search_bank(query: str, limit: int = 20) -> List[Dict[str, Any]]:
        """Text search across the question bank."""
        if not query.strip():
            results = await repo_query(
                "SELECT * FROM question_bank ORDER BY created DESC LIMIT $limit",
                {"limit": limit},
            )
        else:
            results = await repo_query(
                "SELECT * FROM question_bank "
                "WHERE string::contains(string::lowercase(question), string::lowercase($q)) "
                "OR string::contains(string::lowercase(topic), string::lowercase($q)) "
                "LIMIT $limit",
                {"q": query.strip(), "limit": limit},
            )
        return [QuestionPaperService._bank_search_item(r) for r in (results or [])]

    @staticmethod
    async def _load_paper_for_export(paper_id: str) -> dict:
        """Load and validate paper for export."""
        from open_notebook.database.repository import ensure_record_id
        results = await repo_query(
            "SELECT * FROM question_paper WHERE id = $id",
            {"id": ensure_record_id(paper_id)},
        )
        if not results:
            raise HTTPException(status_code=404, detail="Paper not found")
        paper = results[0]
        if paper.get("status") not in ("completed", "needs_manual_review"):
            raise HTTPException(
                status_code=409,
                detail=f"Paper is not yet completed (status: {paper.get('status')})",
            )
        actual_id = str(paper.get("id", paper_id))
        final_paper = paper.get("final_paper") or {}
        stored_count = final_paper.get("question_count", 0)
        section_count = sum(
            len(s.get("questions", [])) for s in final_paper.get("sections", [])
        )
        if stored_count != section_count:
            logger.warning(
                f"Export paper {actual_id}: question_count field ({stored_count}) "
                f"differs from actual section question count ({section_count}). "
                f"Using section count as authoritative."
            )
            final_paper["question_count"] = section_count
        logger.info(
            f"Export loading paper {actual_id}: "
            f"status={paper.get('status')}, accepted_questions={section_count}"
        )
        return paper

    @staticmethod
    async def export_paper_csv(paper_id: str) -> bytes:
        """Export as UTF-8 CSV (legacy, kept for backward compat)."""
        paper = await QuestionPaperService._load_paper_for_export(paper_id)
        final_paper: dict = paper.get("final_paper", {})
        answer_key: List[dict] = paper.get("answer_key", [])
        ak_by_num: Dict[int, dict] = {e["question_number"]: e for e in answer_key if "question_number" in e}

        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow([
            "Q.No", "Section", "Section Ref", "Type", "Answer Type",
            "Target Difficulty", "Validated Difficulty", "Cognitive Score",
            "Question", "Option A", "Option B", "Option C", "Option D", "Option E",
            "Correct Answer", "Explanation",
            "Blind Solver Answer", "Answer Agreement", "Information Sufficient",
            "Arithmetic Consistent", "No Unsupported Claims", "Distractors OK",
            "Option Defensibility",
            "Marks", "Topic", "Sub-topic", "Grade", "Chapter",
        ])

        for section in final_paper.get("sections", []):
            section_name = section.get("section_name", "")
            for q in section.get("questions", []):
                q_num = q.get("question_number", "")
                options: List[str] = q.get("options") or []
                ak_entry = ak_by_num.get(q_num, {})
                writer.writerow([
                    q_num, section_name,
                    q.get("section_ref") or ak_entry.get("section_ref") or "",
                    q.get("type", ""), q.get("answer_type", ""),
                    q.get("target_difficulty") or q.get("difficulty", ""),
                    q.get("validated_cognitive_difficulty") or "",
                    q.get("difficulty_score") or "",
                    q.get("question", ""),
                    options[0] if len(options) > 0 else "",
                    options[1] if len(options) > 1 else "",
                    options[2] if len(options) > 2 else "",
                    options[3] if len(options) > 3 else "",
                    options[4] if len(options) > 4 else "",
                    ak_entry.get("answer", q.get("answer", "")),
                    ak_entry.get("explanation", q.get("explanation", "")),
                    q.get("blind_solver_answer", ""),
                    q.get("answer_agreement", ""),
                    q.get("information_sufficient", ""),
                    q.get("arithmetic_consistent", ""),
                    q.get("no_unsupported_claims", ""),
                    q.get("distractors_ok", ""),
                    json.dumps(q.get("option_defensibility", ""), ensure_ascii=False),
                    q.get("marks", 1), q.get("topic", ""), q.get("sub_topic", ""),
                    q.get("grade", ""), q.get("chapter") or q.get("chapter_title") or "",
                ])

        return ("\ufeff" + output.getvalue()).encode("utf-8")

    @staticmethod
    async def export_paper_xlsx(paper_id: str) -> bytes:
        """Export as Excel (.xlsx) with multiple sheets for QA review."""
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter

        paper = await QuestionPaperService._load_paper_for_export(paper_id)
        final_paper: dict = paper.get("final_paper", {})
        answer_key: List[dict] = paper.get("answer_key", [])
        failed_slots: List[dict] = paper.get("failed_slots") or (final_paper.get("failed_slots") or [])
        audit: dict = paper.get("audit") or {}
        blueprint: dict = paper.get("blueprint") or {}
        status = paper.get("status", "")
        ak_by_num: Dict[int, dict] = {e["question_number"]: e for e in answer_key if "question_number" in e}
        accepted_count = sum(
            len(s.get("questions", [])) for s in final_paper.get("sections", [])
        )

        SCORE_FIELDS = [
            "knowledge", "reasoning", "context", "application",
            "interpretation", "decision_making", "concept_integration", "distractor_quality",
        ]

        bold = Font(bold=True)
        wrap = Alignment(wrap_text=True, vertical="top")
        header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")

        def style_header(ws, col_widths: List[int]):
            for cell in ws[1]:
                cell.font = bold
                cell.fill = header_fill
                cell.alignment = Alignment(wrap_text=True, vertical="center")
            for i, w in enumerate(col_widths, 1):
                ws.column_dimensions[get_column_letter(i)].width = w
            ws.freeze_panes = "A2"
            ws.auto_filter.ref = ws.dimensions

        wb = Workbook()

        # --- Sheet 1: Questions Review ---
        ws = wb.active
        ws.title = "Questions Review"
        q_header = [
            "Q.No", "Grade", "Chapter", "Topic", "Sub-topic",
            "Target Difficulty", "Validated Difficulty", "Total Score",
            "Answer Type", "Question",
            "Option A", "Option B", "Option C", "Option D", "Option E",
            "Correct Answer(s)",
            "Blind Solver Answer",
            "Answer Agreement",
            "Information Sufficient",
            "Arithmetic Consistent",
            "No Unsupported Claims",
            "Distractors OK",
            "Option Defensibility",
            "Attempts", "Status",
        ]
        ws.append(q_header)
        q_widths = [
            6, 6, 20, 20, 20, 14, 14, 8, 14, 50,
            30, 30, 30, 30, 30, 20,
            14, 14, 18, 18, 18, 14, 30,
            8, 12,
        ]

        for section in final_paper.get("sections", []):
            for q in section.get("questions", []):
                q_num = q.get("question_number", "")
                options: List[str] = q.get("options") or []
                ak_entry = ak_by_num.get(q_num, {})
                ws.append([
                    q_num,
                    q.get("grade") or paper.get("grade", ""),
                    q.get("chapter_title") or q.get("chapter") or "",
                    q.get("topic", ""),
                    q.get("sub_topic", ""),
                    q.get("target_difficulty") or "",
                    q.get("validated_cognitive_difficulty") or "",
                    q.get("difficulty_score") or "",
                    q.get("answer_type", ""),
                    q.get("question", ""),
                    options[0] if len(options) > 0 else "",
                    options[1] if len(options) > 1 else "",
                    options[2] if len(options) > 2 else "",
                    options[3] if len(options) > 3 else "",
                    options[4] if len(options) > 4 else "",
                    ak_entry.get("answer", q.get("answer", "")),
                    q.get("blind_solver_answer", ""),
                    q.get("answer_agreement", ""),
                    q.get("information_sufficient", ""),
                    q.get("arithmetic_consistent", ""),
                    q.get("no_unsupported_claims", ""),
                    q.get("distractors_ok", ""),
                    json.dumps(q.get("option_defensibility", ""), ensure_ascii=False),
                    q.get("generation_attempts") or q.get("attempt") or "",
                    q.get("validation_status", ""),
                ])
        style_header(ws, q_widths)
        for row in ws.iter_rows(min_row=2, max_col=len(q_header)):
            for cell in row:
                cell.alignment = wrap

        # --- Sheet 2: Cognitive Validation ---
        ws2 = wb.create_sheet("Cognitive Validation")
        cv_header = ["Q.No", "Chapter", "Target Difficulty", "Validated Difficulty"] + \
            [s.replace("_", " ").title() for s in SCORE_FIELDS] + \
            ["Total Score", "Validation Status", "Validation Reasons"]
        ws2.append(cv_header)
        cv_widths = [6, 20, 14, 14] + [12] * 8 + [8, 12, 60]

        for section in final_paper.get("sections", []):
            for q in section.get("questions", []):
                scores = q.get("difficulty_scores") or {}
                ws2.append([
                    q.get("question_number", ""),
                    q.get("chapter_title") or q.get("chapter") or "",
                    q.get("target_difficulty") or "",
                    q.get("validated_cognitive_difficulty") or "",
                ] + [scores.get(f, "") for f in SCORE_FIELDS] + [
                    q.get("difficulty_score") or "",
                    q.get("validation_status", ""),
                    "; ".join(q.get("validation_reasons") or []),
                ])
        style_header(ws2, cv_widths)

        # --- Sheet 3: Answer Key ---
        ws3 = wb.create_sheet("Answer Key")
        ws3.append(["Q.No", "Question", "Correct Answer(s)", "Explanation"])
        ak_widths = [6, 50, 20, 60]
        for item in answer_key:
            ws3.append([
                item.get("question_number", ""),
                item.get("question", ""),
                item.get("answer", ""),
                item.get("explanation", ""),
            ])
        style_header(ws3, ak_widths)
        for row in ws3.iter_rows(min_row=2, max_col=4):
            for cell in row:
                cell.alignment = wrap

        # --- Sheet 4: Audit Summary ---
        ws4 = wb.create_sheet("Audit Summary")
        ws4.append(["Field", "Value"])
        ws4.column_dimensions["A"].width = 35
        ws4.column_dimensions["B"].width = 60
        for cell in ws4[1]:
            cell.font = bold
            cell.fill = header_fill
        ws4.append(["Status", status.upper().replace("_", " ")])
        ws4.append(["Audit OK", str(audit.get("ok", "N/A"))])
        ws4.append(["Topic", paper.get("topic", "")])
        ws4.append(["Grade", paper.get("grade", "")])
        ws4.append(["Subject", paper.get("subject", "")])
        ws4.append(["Expected Questions", paper.get("target_marks", "")])
        ws4.append(["Accepted Questions / Actual Marks", accepted_count])
        ws4.append(["Failed/Manual Review Slots", len(failed_slots)])
        ws4.append([""])
        ws4.append(["DIFFICULTY DISTRIBUTION (TARGET)", ""])
        chapters_bp = blueprint.get("chapter_difficulty") or blueprint.get("chapters") or {}
        diff_totals: Dict[str, int] = {"easy": 0, "medium": 0, "difficult": 0}
        for ch_counts in chapters_bp.values():
            if isinstance(ch_counts, dict):
                for d, v in ch_counts.items():
                    diff_totals[d] = diff_totals.get(d, 0) + int(v)
        for d, count in diff_totals.items():
            ws4.append([f"  {d.capitalize()}", count])
        ws4.append([""])
        ws4.append(["DIFFICULTY DISTRIBUTION (ACTUAL ACCEPTED)", ""])
        actual_diff: Dict[str, int] = {"easy": 0, "medium": 0, "difficult": 0}
        for section in final_paper.get("sections", []):
            for q in section.get("questions", []):
                vd = q.get("validated_cognitive_difficulty") or q.get("target_difficulty") or ""
                if vd in actual_diff:
                    actual_diff[vd] += 1
        for d, count in actual_diff.items():
            ws4.append([f"  {d.capitalize()}", count])
        ws4.append([""])
        ws4.append(["AUDIT ERRORS", ""])
        for err in audit.get("errors") or []:
            ws4.append(["", err])

        # --- Sheet 5: Failed Slots ---
        ws5 = wb.create_sheet("Failed Slots")
        fs_header = ["Q.No", "Chapter", "Target Difficulty", "Answer Type", "Attempts", "Rejection Reasons"]
        ws5.append(fs_header)
        fs_widths = [6, 20, 14, 14, 8, 80]
        for slot in failed_slots:
            ws5.append([
                slot.get("question_number", ""),
                slot.get("chapter_title") or slot.get("chapter", ""),
                slot.get("target_difficulty", ""),
                slot.get("answer_type", ""),
                slot.get("attempts") or slot.get("generation_attempts") or "",
                "; ".join(slot.get("validation_reasons") or slot.get("reasons") or []),
            ])
        style_header(ws5, fs_widths)
        for row in ws5.iter_rows(min_row=2, max_col=len(fs_header)):
            for cell in row:
                cell.alignment = wrap

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    @staticmethod
    async def export_paper_docx(paper_id: str) -> bytes:
        """Export as Word (.docx) — clean academic paper with short answer key."""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        paper = await QuestionPaperService._load_paper_for_export(paper_id)
        final_paper: dict = paper.get("final_paper", {})
        answer_key: List[dict] = paper.get("answer_key", [])
        status = paper.get("status", "")
        target_marks = paper.get("target_marks", 0)
        accepted_count = sum(
            len(s.get("questions", [])) for s in final_paper.get("sections", [])
        )

        doc = Document()

        # Draft label
        if status == "needs_manual_review":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("DRAFT \u2014 NEEDS MANUAL REVIEW")
            run.bold = True
            run.font.size = Pt(14)

        # Title
        title = doc.add_heading(f"Question Paper \u2014 {paper.get('topic', '')}", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parts = []
        if paper.get("grade"):
            parts.append(f"Grade: {paper['grade']}")
        if paper.get("subject"):
            parts.append(f"Subject: {paper['subject']}")
        parts.append(f"Expected Questions: {target_marks}")
        parts.append(f"Accepted Questions: {accepted_count}")
        parts.append(f"Total Marks: {accepted_count}")
        meta.add_run("   |   ".join(parts))

        doc.add_paragraph()

        # Questions — clean format
        for section in final_paper.get("sections", []):
            for q in section.get("questions", []):
                q_num = q.get("question_number", "")
                marks = q.get("marks", 1)
                p = doc.add_paragraph()
                p.add_run(f"{q_num}. ").bold = True
                p.add_run(q.get("question", ""))
                p.add_run(f"  [{marks} mark{'s' if marks != 1 else ''}]").italic = True
                options: List[str] = q.get("options") or []
                for i, opt in enumerate(options):
                    doc.add_paragraph(f"     {chr(65 + i)}.  {opt}")

        # Page break before answer key
        doc.add_page_break()
        doc.add_heading("Answer Key", level=1)
        for item in answer_key:
            p = doc.add_paragraph()
            p.add_run(f"{item.get('question_number', '')}. ").bold = True
            p.add_run(str(item.get("answer", "")))

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    async def export_paper_txt(paper_id: str) -> bytes:
        """Export as plain text with correct metadata."""
        paper = await QuestionPaperService._load_paper_for_export(paper_id)
        final_paper: dict = paper.get("final_paper", {})
        answer_key: List[dict] = paper.get("answer_key", [])
        status = paper.get("status", "")
        blueprint: dict = paper.get("blueprint") or {}
        accepted_count = sum(
            len(s.get("questions", [])) for s in final_paper.get("sections", [])
        )

        lines: List[str] = []

        if status == "needs_manual_review":
            lines.append("*** DRAFT — NEEDS MANUAL REVIEW ***")
            lines.append("")

        lines.append(f"QUESTION PAPER — {paper.get('topic', '').upper()}")
        if paper.get("grade"):
            lines.append(f"Grade: {paper['grade']}")
        if paper.get("subject"):
            lines.append(f"Subject: {paper['subject']}")
        lines.append(f"Expected Questions/Marks: {paper.get('target_marks', '')}")
        lines.append(f"Accepted Questions/Actual Marks: {accepted_count}")

        # Difficulty distribution
        chapters_bp = blueprint.get("chapter_difficulty") or blueprint.get("chapters") or {}
        diff_totals: Dict[str, int] = {"easy": 0, "medium": 0, "difficult": 0}
        for ch_counts in chapters_bp.values():
            if isinstance(ch_counts, dict):
                for d, v in ch_counts.items():
                    diff_totals[d] = diff_totals.get(d, 0) + int(v)
        if any(diff_totals.values()):
            lines.append(f"Target Difficulty: Easy={diff_totals['easy']}, Medium={diff_totals['medium']}, Difficult={diff_totals['difficult']}")

        lines.append("=" * 60)
        lines.append("")

        for section in final_paper.get("sections", []):
            lines.append(f"SECTION: {section.get('section_name', '').upper()}")
            lines.append("-" * 40)
            for q in section.get("questions", []):
                lines.append(f"{q.get('question_number', '')}. {q.get('question', '')} [{q.get('marks', 1)} mark(s)]")
                for i, opt in enumerate(q.get("options") or []):
                    lines.append(f"   {chr(65 + i)}. {opt}")
                lines.append("")

        lines.append("=" * 60)
        lines.append("ANSWER KEY")
        lines.append("=" * 60)
        for item in answer_key:
            lines.append(f"{item.get('question_number', '')}. {item.get('answer', '')}")
            if item.get("explanation"):
                lines.append(f"   Explanation: {item['explanation']}")
            lines.append("")

        return "\n".join(lines).encode("utf-8")

    @staticmethod
    async def delete_bank_question(question_id: str) -> None:
        """Remove a question from the bank."""
        from open_notebook.database.repository import ensure_record_id
        try:
            await repo_delete(ensure_record_id(question_id))
        except Exception as e:
            logger.error(f"Failed to delete question {question_id}: {e}")
            raise HTTPException(status_code=500, detail="Failed to delete question")

    # ------------------------------------------------------------------
    # Book upload helpers
    # ------------------------------------------------------------------

    _FRONT_MATTER_TITLES = frozenset({
        "preface", "foreword", "introduction", "contents", "table of contents",
        "acknowledgements", "acknowledgments", "copyright", "dedication",
        "about the author", "about the book", "index",
    })

    @staticmethod
    async def get_book_content(
        book_id: str,
        selected_chapters: Optional[List[int]],
        skip_front_matter: bool = True,
    ) -> str:
        """Fetch book text for selected chapters (or full text if none selected).

        Returns annotated text with === SECTION === markers so the LLM can
        populate section_ref on generated questions.
        """
        from open_notebook.database.repository import ensure_record_id
        results = await repo_query(
            "SELECT * FROM question_book WHERE id = $id",
            {"id": ensure_record_id(book_id)},
        )
        if not results:
            raise HTTPException(status_code=404, detail="Book not found")
        book = results[0]
        full_text: str = book.get("full_text", "")
        chapters: List[dict] = book.get("chapters", [])

        if not chapters:
            return full_text

        # Determine which chapters to include
        chapter_indices: List[int]
        if selected_chapters is None:
            chapter_indices = list(range(len(chapters)))
        else:
            chapter_indices = selected_chapters

        # Auto-skip front-matter chapters when no explicit selection was made
        if skip_front_matter and selected_chapters is None:
            chapter_indices = [
                idx for idx in chapter_indices
                if chapters[idx].get("title", "").strip().lower()
                not in QuestionPaperService._FRONT_MATTER_TITLES
            ]
            if not chapter_indices:
                # Nothing left — fall back to all chapters
                chapter_indices = list(range(len(chapters)))

        # Build annotated sections with === markers so section_ref can be set
        parts = []
        for idx in chapter_indices:
            if 0 <= idx < len(chapters):
                ch = chapters[idx]
                start = ch.get("start_char", 0)
                end = ch.get("end_char", len(full_text))
                label = ch.get("title", f"Section {idx + 1}")
                section_text = full_text[start:end]
                parts.append(f"=== SECTION {idx + 1}: {label} ===\n{section_text}")

        return "\n\n".join(parts) if parts else full_text

    @staticmethod
    async def get_book_chapters(
        book_id: str,
        selected_chapters: Optional[List[int]],
        skip_front_matter: bool = True,
    ) -> List[Dict[str, Any]]:
        """Return selected chapters as {title, text} preserving full chapter content.

        Does not truncate to a character budget. Callers chunk per chapter later.
        """
        from open_notebook.database.repository import ensure_record_id
        results = await repo_query(
            "SELECT * FROM question_book WHERE id = $id",
            {"id": ensure_record_id(book_id)},
        )
        if not results:
            raise HTTPException(status_code=404, detail="Book not found")
        book = results[0]
        full_text: str = book.get("full_text", "")
        chapters: List[dict] = book.get("chapters", [])

        if not chapters:
            return [{"title": book.get("title") or "Full material", "text": full_text}] if full_text else []

        if selected_chapters is None:
            chapter_indices = list(range(len(chapters)))
        else:
            chapter_indices = selected_chapters

        if skip_front_matter and selected_chapters is None:
            chapter_indices = [
                idx for idx in chapter_indices
                if chapters[idx].get("title", "").strip().lower()
                not in QuestionPaperService._FRONT_MATTER_TITLES
            ]
            if not chapter_indices:
                chapter_indices = list(range(len(chapters)))

        parsed: List[Dict[str, Any]] = []
        for idx in chapter_indices:
            if 0 <= idx < len(chapters):
                ch = chapters[idx]
                start = ch.get("start_char", 0)
                end = ch.get("end_char", len(full_text))
                parsed.append({
                    "index": idx,
                    "title": ch.get("title", f"Section {idx + 1}"),
                    "text": full_text[start:end],
                })
        return parsed

    @staticmethod
    async def create_and_submit_bank_batch(
        request: GenerateBankBatchRequest,
    ) -> GenerateBankBatchResponse:
        """Create a question_bank_batch record and submit the async generation job."""
        from open_notebook.graphs.question_paper_blueprint import bank_batch_from_dict

        try:
            import commands.question_paper_commands  # noqa: F401
        except ImportError as e:
            logger.warning(f"Could not import question paper commands module: {e}")

        if request.max_slot_attempts < 1 or request.max_slot_attempts > 10:
            raise HTTPException(status_code=400, detail="max_slot_attempts must be between 1 and 10")
        if request.max_refill_attempts < 0 or request.max_refill_attempts > 10:
            raise HTTPException(status_code=400, detail="max_refill_attempts must be between 0 and 10")
        if request.slot_concurrency < 1 or request.slot_concurrency > 8:
            raise HTTPException(status_code=400, detail="slot_concurrency must be between 1 and 8")
        if request.total_questions < 1 or request.total_questions > 200:
            raise HTTPException(
                status_code=400,
                detail="total_questions must be between 1 and 200",
            )

        blueprint_dict = {
            "book_id": request.book_id.strip(),
            "grade": request.grade.strip(),
            "subject": request.subject.strip(),
            "chapter": request.chapter,
            "difficulty": request.difficulty,
            "total_questions": request.total_questions,
            "single_correct": request.single_correct,
            "multiple_correct": request.multiple_correct,
            "language": request.language,
        }
        try:
            blueprint = bank_batch_from_dict(blueprint_dict)
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))

        chapter_index = int(blueprint.chapter) - 1
        if chapter_index < 0:
            raise HTTPException(status_code=400, detail="chapter must be >= 1")

        book_chapters = await QuestionPaperService.get_book_chapters(
            blueprint.book_id,
            selected_chapters=[chapter_index],
            skip_front_matter=False,
        )
        if not book_chapters:
            raise HTTPException(
                status_code=400,
                detail="No chapter content could be loaded for the selected chapter.",
            )
        book_content = await QuestionPaperService.get_book_content(
            blueprint.book_id,
            selected_chapters=[chapter_index],
            skip_front_matter=False,
        )

        batch_data = {
            "book_id": blueprint.book_id,
            "grade": blueprint.grade,
            "subject": blueprint.subject,
            "chapter": blueprint.chapter,
            "difficulty": blueprint.difficulty,
            "total_questions": blueprint.total_questions,
            "single_correct": blueprint.single_correct,
            "multiple_correct": blueprint.multiple_correct,
            "language": blueprint.language,
            "status": "pending",
            "requested": blueprint.total_questions,
            "blueprint": blueprint_dict,
        }
        raw_batch = await repo_create("question_bank_batch", batch_data)
        if not raw_batch:
            raise HTTPException(status_code=500, detail="Failed to create bank batch record")
        batch_record = raw_batch[0] if isinstance(raw_batch, list) else raw_batch
        batch_id = str(batch_record["id"])

        try:
            cmd_id = submit_command(
                "open_notebook",
                "generate_question_bank_batch",
                {
                    "batch_id": batch_id,
                    "book_id": blueprint.book_id,
                    "grade": blueprint.grade,
                    "subject": blueprint.subject,
                    "chapter": blueprint.chapter,
                    "difficulty": blueprint.difficulty,
                    "total_questions": blueprint.total_questions,
                    "single_correct": blueprint.single_correct,
                    "multiple_correct": blueprint.multiple_correct,
                    "language": blueprint.language,
                    "book_content": book_content,
                    "book_chapters": book_chapters,
                    "curriculum_objectives": request.curriculum_objectives,
                    "generator_model": request.generator_model,
                    "reviewer_model": request.reviewer_model,
                    "max_slot_attempts": request.max_slot_attempts,
                    "max_refill_attempts": request.max_refill_attempts,
                    "slot_concurrency": request.slot_concurrency,
                },
            )
        except Exception as e:
            logger.error(f"Failed to submit bank batch job: {e}")
            raise HTTPException(status_code=500, detail="Failed to submit bank batch generation job")

        cmd_id_str = str(cmd_id)
        from open_notebook.database.repository import repo_query

        try:
            await repo_query(
                f"UPDATE {batch_id} SET command = {cmd_id_str}",
                {},
            )
        except Exception as e:
            logger.warning(f"Could not store command reference on batch {batch_id}: {e}")

        logger.info(f"Submitted bank batch job {cmd_id_str} for batch {batch_id}")

        return GenerateBankBatchResponse(
            job_id=cmd_id_str,
            batch_id=batch_id,
            status="submitted",
            message=(
                f"Question bank batch started: {blueprint.total_questions} "
                f"{blueprint.difficulty} question(s) for chapter {blueprint.chapter}"
            ),
            requested=blueprint.total_questions,
        )

    @staticmethod
    async def get_bank_batch_status(batch_id: str) -> Dict[str, Any]:
        """Get status of a question bank batch job."""
        from open_notebook.database.repository import ensure_record_id

        results = await repo_query(
            "SELECT * FROM question_bank_batch WHERE id = $id",
            {"id": ensure_record_id(batch_id)},
        )
        if not results:
            raise HTTPException(status_code=404, detail="Bank batch not found")

        batch = results[0]
        batch_status = batch.get("status", "unknown")
        job_status = None

        command_id = batch.get("command")
        if command_id:
            try:
                cmd = await get_command_status(str(command_id))
                if cmd:
                    job_status = cmd.status
                    if cmd.status in ("failed", "error") and batch_status == "running":
                        batch_status = "failed"
            except Exception as e:
                logger.warning(f"Failed to get command status for {command_id}: {e}")

        return {
            "batch_id": str(batch.get("id", batch_id)),
            "status": batch_status,
            "job_status": job_status,
            "book_id": batch.get("book_id"),
            "grade": batch.get("grade"),
            "subject": batch.get("subject"),
            "chapter": batch.get("chapter"),
            "difficulty": batch.get("difficulty"),
            "requested": batch.get("requested"),
            "accepted": batch.get("accepted"),
            "failed": batch.get("failed"),
            "error_message": batch.get("error_message"),
            "created": str(batch.get("created", "")),
        }

    @staticmethod
    async def get_bank_batch_result(batch_id: str) -> Dict[str, Any]:
        """Fetch a completed or partially completed bank batch result."""
        from open_notebook.database.repository import ensure_record_id

        results = await repo_query(
            "SELECT * FROM question_bank_batch WHERE id = $id",
            {"id": ensure_record_id(batch_id)},
        )
        if not results:
            raise HTTPException(status_code=404, detail="Bank batch not found")

        batch = results[0]
        status = batch.get("status", "")
        if status not in ("completed", "completed_partial", "failed"):
            raise HTTPException(
                status_code=409,
                detail=f"Bank batch is not yet finished (status: {status})",
            )

        saved_ids = batch.get("saved_question_ids") or []
        questions: List[dict] = []
        if saved_ids:
            from open_notebook.database.repository import ensure_record_id

            q_rows = await repo_query(
                "SELECT * FROM question_bank WHERE id IN $ids",
                {"ids": [ensure_record_id(i) for i in saved_ids]},
            )
            questions = q_rows or []

        return {
            "batch_id": str(batch.get("id", batch_id)),
            "status": status,
            "book_id": batch.get("book_id"),
            "grade": batch.get("grade"),
            "subject": batch.get("subject"),
            "chapter": batch.get("chapter"),
            "difficulty": batch.get("difficulty"),
            "requested": batch.get("requested"),
            "accepted": batch.get("accepted"),
            "failed": batch.get("failed"),
            "audit": batch.get("audit"),
            "failure_summary": batch.get("failure_summary"),
            "failed_slots": batch.get("failed_slots") or [],
            "rejected_attempts": batch.get("rejected_attempts") or [],
            "saved_question_ids": saved_ids,
            "questions": questions,
            "error_message": batch.get("error_message"),
        }


class BookService:
    """Handles book upload, text extraction, and chapter detection."""

    @staticmethod
    def build_display_name(
        book_name: Optional[str],
        year: Optional[str],
        grade: Optional[str],
        override: Optional[str] = None,
    ) -> Optional[str]:
        if override and str(override).strip():
            return str(override).strip()
        name = (book_name or "").strip()
        yr = str(year).strip() if year is not None else ""
        gr = str(grade).strip() if grade is not None else ""
        if name and yr and gr:
            return f"{name} {yr} - Grade {gr}"
        return None

    @staticmethod
    def serialize_book(record: Dict[str, Any], *, include_chapters: bool = False) -> Dict[str, Any]:
        book_name = (record.get("book_name") or "").strip() or None
        year = record.get("year")
        year_str = str(year).strip() if year is not None and str(year).strip() else None
        grade = (record.get("grade") or "").strip() or None
        detected_grade = (record.get("detected_grade") or "").strip() or None
        grade_for_filter = grade or detected_grade
        display_override = (record.get("display_name") or "").strip() or None
        display_name = BookService.build_display_name(
            book_name, year_str, grade_for_filter, display_override
        )
        title = (record.get("title") or "").strip() or None
        missing_fields: List[str] = []
        if not book_name:
            missing_fields.append("Book Name")
        if not year_str:
            missing_fields.append("Year")
        payload: Dict[str, Any] = {
            "book_id": str(record.get("id", "")),
            "book_id": str(record.get("id", "")),
            "book_name": book_name,
            "year": year_str,
            "grade": grade_for_filter,
            "subject": (record.get("subject") or "").strip() or None,
            "edition": (record.get("edition") or "").strip() or None,
            "display_name": display_name or title or str(record.get("id", "")),
            "title": title,
            "detected_grade": detected_grade,
            "chapter_count": len(record.get("chapters") or []),
            "missing_fields": missing_fields,
            "metadata_complete": len(missing_fields) == 0,
        }
        if include_chapters:
            chapters = record.get("chapters") or []
            mapped = []
            total_chars = 0
            for ch in chapters:
                start = int(ch.get("start_char") or 0)
                end = int(ch.get("end_char") or 0)
                char_count = max(0, end - start)
                total_chars += char_count
                mapped.append({
                    "index": ch.get("index", 0),
                    "title": ch.get("title") or "",
                    "preview": ch.get("preview") or "",
                    "char_count": char_count,
                })
            payload["chapters"] = mapped
            payload["total_chars"] = total_chars
            payload["total_chars"] = total_chars
        return payload

    @staticmethod
    async def list_books() -> List[Dict[str, Any]]:
        results = await repo_query(
            "SELECT id, title, book_name, year, grade, subject, edition, display_name, "
            "detected_grade, chapters, created FROM question_book ORDER BY created DESC"
        )
        return [BookService.serialize_book(row, include_chapters=False) for row in (results or [])]

    @staticmethod
    async def update_metadata(
        book_id: str,
        *,
        book_name: str,
        year: str,
        grade: str,
        subject: Optional[str] = None,
        edition: Optional[str] = None,
        display_name: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Update library labels only. Preserves id, full_text, chapters, and file_path."""
        from open_notebook.database.repository import ensure_record_id

        record_id = ensure_record_id(book_id)
        existing = await repo_query(
            "SELECT id FROM question_book WHERE id = $id",
            {"id": record_id},
        )
        if not existing:
            raise HTTPException(status_code=404, detail="Book not found")
        original_id = str(existing[0].get("id", ""))

        stored_name = (book_name or "").strip() or None
        stored_year = (year or "").strip() or None
        stored_grade = (grade or "").strip() or None
        if not stored_name or not stored_year or not stored_grade:
            raise HTTPException(
                status_code=400,
                detail="Book Name, Year, and Grade are required.",
            )
        stored_subject = (subject or "").strip() or None
        stored_edition = (edition or "").strip() or None
        built_display = BookService.build_display_name(
            stored_name, stored_year, stored_grade, display_name
        )
        await repo_update(
            "question_book",
            str(record_id),
            {
                "book_name": stored_name,
                "year": stored_year,
                "grade": stored_grade,
                "subject": stored_subject,
                "edition": stored_edition,
                "display_name": built_display,
            },
        )
        updated = await repo_query(
            "SELECT id, title, book_name, year, grade, subject, edition, display_name, "
            "detected_grade, chapters, created FROM question_book WHERE id = $id",
            {"id": record_id},
        )
        if not updated:
            raise HTTPException(status_code=500, detail="Book metadata update did not return the record")
        if str(updated[0].get("id", "")) != original_id:
            raise HTTPException(status_code=500, detail="Book id changed unexpectedly; update aborted")
        return BookService.serialize_book(updated[0], include_chapters=True)

    @staticmethod
    async def upload_and_extract(
        upload_file: UploadFile,
        *,
        book_name: Optional[str] = None,
        year: Optional[str] = None,
        grade: Optional[str] = None,
        subject: Optional[str] = None,
        edition: Optional[str] = None,
        display_name: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Save the file, extract full text via content_core, detect chapters.
        Returns book record with chapter list.
        """
        from content_core import extract_content

        from open_notebook.config import UPLOADS_FOLDER

        if not upload_file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")

        # 1. Save to uploads folder
        safe_name = re.sub(r"[^\w\-_\. ]", "_", upload_file.filename)
        dest = os.path.join(UPLOADS_FOLDER, safe_name)
        # Make unique
        base, ext = os.path.splitext(dest)
        counter = 1
        while os.path.exists(dest):
            dest = f"{base}_{counter}{ext}"
            counter += 1

        try:
            content = await upload_file.read()
            with open(dest, "wb") as f:
                f.write(content)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to save file: {e}")

        # 2. Extract text — content_core expects a state dict and is async
        try:
            from content_core import extract_content
            content_state = {
                "file_path": dest,
                "document_engine": "auto",
                "output_format": "markdown",
                "delete_source": False,
            }
            processed = await extract_content(content_state)
            # extract_content returns a ProcessSourceState object; .content holds the text
            full_text: str = getattr(processed, "content", "") or ""
            if not full_text:
                # fallback: try dict-style access
                full_text = processed.get("content", "") if isinstance(processed, dict) else ""
        except Exception as e:
            logger.error(f"content_core extraction failed: {e}")
            # Plain-text fallback for .txt / .md files
            try:
                with open(dest, "r", encoding="utf-8", errors="ignore") as f:
                    full_text = f.read()
            except Exception:
                raise HTTPException(status_code=422, detail=f"Failed to extract text: {e}")

        if not full_text.strip():
            raise HTTPException(status_code=422, detail="Could not extract any text from the file")

        # 3. Detect chapters
        chapters = BookService._detect_chapters(full_text)

        # 4. Detect grade from content
        detected_grade = BookService._detect_grade(full_text, upload_file.filename or "")

        # 5. Store in DB (new metadata is optional for backward compatibility)
        file_stem = os.path.splitext(upload_file.filename)[0]
        stored_grade = (grade or "").strip() or None
        stored_name = (book_name or "").strip() or None
        stored_year = (year or "").strip() or None
        stored_subject = (subject or "").strip() or None
        stored_edition = (edition or "").strip() or None
        built_display = BookService.build_display_name(
            stored_name, stored_year, stored_grade, display_name
        )
        title = built_display or stored_name or file_stem
        raw = await repo_create("question_book", {
            "title": title,
            "file_path": dest,
            "full_text": full_text,
            "chapters": chapters,
            "detected_grade": detected_grade,
            "book_name": stored_name,
            "year": stored_year,
            "grade": stored_grade,
            "subject": stored_subject,
            "edition": stored_edition,
            "display_name": built_display or title,
        })
        record = raw[0] if isinstance(raw, list) else raw
        serialized = BookService.serialize_book(
            {**record, "chapters": chapters, "full_text": full_text},
            include_chapters=True,
        )
        serialized["total_chars"] = len(full_text)
        return serialized

    @staticmethod
    def _detect_chapters(text: str) -> List[Dict[str, Any]]:
        """
        Hierarchical chapter detection tried in order of specificity.

        Strategy 1 — explicit keyword  ("Chapter 1 / Chapter I"):       min 3000 chars
        Strategy 2 — pipe format       ("1 | Title"):                   min 2000 chars
        Strategy 3 — Markdown H1       ("# Title"):                     min 2000 chars
        Strategy 4 — isolated numbered ("1. Title" preceded by blank):  min 3000 chars
        Strategy 5 — Markdown H2       ("## Title"):                    min 1000 chars
        Strategy 6 — subsection num    ("1.1 Title"):                   min  500 chars
        Fallback   — equal-size chunks

        Each strategy is tried independently. The first one that yields >= 2 chapters wins.
        Numbered list items (e.g. "5. The British Pound Sterling:") are excluded from
        strategy 4 by requiring a blank line before the heading.
        """

        def _has_blank_line_before(pos: int) -> bool:
            """Return True if the character before pos is preceded by a blank line."""
            before = text[:pos]
            last_nl = before.rfind("\n")
            if last_nl == -1:
                return True  # start of file
            second_nl = before.rfind("\n", 0, last_nl)
            if second_nl == -1:
                return pos < 10  # very near start of file
            between = before[second_nl + 1 : last_nl].strip()
            return len(between) == 0

        def _filter_matches(matches: list, min_chars: int, require_blank_before: bool = False) -> list:
            kept = []
            for i, m in enumerate(matches):
                if require_blank_before and not _has_blank_line_before(m.start()):
                    continue
                next_start = matches[i + 1].start() if i + 1 < len(matches) else len(text)
                if (next_start - m.end()) >= min_chars:
                    kept.append(m)
            return kept

        def _build_chapters(matches: list) -> List[Dict[str, Any]]:
            chapters = []
            for i, match in enumerate(matches):
                start = match.start()
                end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
                chapter_text = text[start:end].strip()
                heading = next(g for g in match.groups() if g is not None).strip()
                preview = chapter_text[:200].replace("\n", " ").strip()
                chapters.append({
                    "index": i,
                    "title": heading[:120],
                    "start_char": start,
                    "end_char": end,
                    "preview": preview,
                })
            return chapters

        def _try(pattern_str: str, flags: int, min_chars: int, require_blank: bool = False) -> List[Dict[str, Any]]:
            matches = list(re.finditer(pattern_str, text, flags))
            filtered = _filter_matches(matches, min_chars, require_blank_before=require_blank)
            if len(filtered) >= 2:
                return _build_chapters(filtered)
            return []

        # Strategy 1: explicit chapter keyword — most reliable
        # Use [ \t] (not \s) in the separator so newlines don't get consumed into the title.
        # Require the number part to be purely digits or roman numerals (not mixed like "80D").
        result = _try(
            r"^((?:chapter|unit|module|lesson)\s+(?:\d{1,3}|[ivxlcdmIVXLCDM]+)(?![A-Za-z])[.\ \t:—\-]*[^\n]{0,80})",
            re.MULTILINE | re.IGNORECASE,
            min_chars=3000,
        )
        if result:
            return result

        # Strategy 2: pipe-delimited "1 | Title" — common in textbooks/workbooks
        result = _try(
            r"^(\d{1,2}\s*[|｜]\s+[A-Z][^\n]{2,60})$",
            re.MULTILINE,
            min_chars=2000,
        )
        if result:
            return result

        # Strategy 3: Markdown H1
        result = _try(
            r"^(#\s+[^\n]{3,80})",
            re.MULTILINE,
            min_chars=2000,
        )
        if result:
            return result

        # Strategy 4: isolated numbered heading "1. Title" — REQUIRE blank line before
        # to avoid matching items inside numbered lists.
        result = _try(
            r"^(\d{1,2}\.\s+[A-Z][^\n]{5,60})$",
            re.MULTILINE,
            min_chars=3000,
            require_blank=True,
        )
        if result:
            return result

        # Strategy 5: Markdown H2
        result = _try(
            r"^(#{1,2}\s+[^\n]{3,80})",
            re.MULTILINE,
            min_chars=1000,
        )
        if result:
            return result

        # Strategy 6: subsection numbering "1.1 Title"
        result = _try(
            r"^(\d{1,2}\.\d+\s+[A-Z][^\n]{5,60})$",
            re.MULTILINE,
            min_chars=500,
        )
        if result:
            return result

        # Fallback: equal-size chunks
        return BookService._chunk_fallback(text)

    @staticmethod
    def _detect_grade(text: str, filename: str) -> Optional[str]:
        """Best-effort grade detection from filename and first ~2000 chars of content."""
        combined = f"{filename}\n{text[:2000]}"
        patterns = [
            r"\bgrade\s*[-:.]?\s*(\d{1,2})\b",
            r"\bclass\s*[-:.]?\s*(\d{1,2})\b",
            r"\bstd\.?\s*(\d{1,2})\b",
            r"\bstandard\s*[-:.]?\s*(\d{1,2})\b",
        ]
        for pat in patterns:
            match = re.search(pat, combined, re.IGNORECASE)
            if match:
                num = int(match.group(1))
                if 1 <= num <= 12:
                    return str(num)
        return None

    @staticmethod
    def _chunk_fallback(text: str, chunk_size: int = 5000) -> List[Dict[str, Any]]:
        """Split text into equal chunks when no chapter structure is detected."""
        chunks = []
        for i, start in enumerate(range(0, len(text), chunk_size)):
            end = min(start + chunk_size, len(text))
            snippet = text[start:end].strip()
            chunks.append({
                "index": i,
                "title": f"Part {i + 1}",
                "start_char": start,
                "end_char": end,
                "preview": snippet[:200].replace("\n", " "),
            })
        return chunks
